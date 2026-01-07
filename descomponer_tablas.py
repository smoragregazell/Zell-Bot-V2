from docx import Document

DOCX_PATH = r"C:\Users\Santiago Moragrega\Zell_Bot_V2\tablas proceso de atención.docx"

def cell_merge_info(cell):
    """
    Lee metadata de merges desde el XML del <w:tcPr>.
    - gridSpan = merge horizontal (colspan)
    - vMerge   = merge vertical (rowspan-ish)
    """
    tc = cell._tc
    tcPr = tc.tcPr
    info = {}
    if tcPr is None:
        return info

    # horizontal merge
    gs = tcPr.gridSpan
    if gs is not None and gs.val is not None:
        info["gridSpan"] = int(gs.val)

    # vertical merge
    vm = tcPr.vMerge
    if vm is not None:
        # vm.val puede ser None (continuación) o "restart"
        info["vMerge"] = vm.val if vm.val is not None else "continue"

    return info

def table_to_matrix(tbl):
    return [[cell.text.strip() for cell in row.cells] for row in tbl.rows]

def dump_docx_tables(path: str, sample_rows: int = None, max_merge_markers: int = None):
    doc = Document(path)
    print(f"Doc: {path}")
    print(f"Tablas: {len(doc.tables)}\n")

    for ti, tbl in enumerate(doc.tables):
        mat = table_to_matrix(tbl)
        rows = len(mat)
        cols = max((len(r) for r in mat), default=0)

        print("=" * 80)
        print(f"Tabla {ti}: {rows} filas x {cols} columnas")

        if rows > 0:
            headers = mat[0]
            print("Headers:", " | ".join(headers))

        # Mostrar todas las filas si sample_rows es None, sino mostrar solo las primeras N
        if sample_rows is None:
            print("\nTodas las filas:")
            for r in mat:
                print(" - " + " | ".join(r))
        else:
            print(f"\nMuestra (primeras {sample_rows} filas):")
            for r in mat[:sample_rows]:
                print(" - " + " | ".join(r))

        # merges
        merges = []
        for r_i, row in enumerate(tbl.rows):
            for c_i, cell in enumerate(row.cells):
                info = cell_merge_info(cell)
                if info:
                    merges.append((r_i, c_i, info))

        if merges:
            print("\nMerges detectados (fila, col, info):")
            if max_merge_markers is None:
                # Mostrar todos los merges
                for item in merges:
                    print(" -", item)
            else:
                # Mostrar solo los primeros N merges
                for item in merges[:max_merge_markers]:
                    print(" -", item)
                if len(merges) > max_merge_markers:
                    print(f" ... ({len(merges)} en total)")
        else:
            print("\nMerges detectados: ninguno")

        print()

if __name__ == "__main__":
    # Mostrar todas las filas y todos los merges
    dump_docx_tables(DOCX_PATH, sample_rows=None, max_merge_markers=None)
